import streamlit as st
import pandas as pd
import numpy as np
import io
from reportlab.lib.pagesizes import letter, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Inches

def categorize_attendance(percentage):
    percentage = percentage * 100
    if percentage >= 75:
        return "≥75%"
    elif 65 <= percentage < 75:
        return "≥65% - <75%"
    elif 50 <= percentage < 65:
        return "≥50% - <65%"
    else:
        return "<50%"

def categorize_attendance_per(percentage):
    if percentage >= 75:
        return "75% and above"
    elif 65 <= percentage < 75:
        return "65–74%"
    elif 50 <= percentage < 65:
        return "50–64%"
    else:
        return "0–49%"

def create_pdf(sheet_name, mode, report_df, report1_df):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    elements = []
    styles = getSampleStyleSheet()
    
    elements.append(Paragraph(f"Student FA Marks Report - {sheet_name} ({mode} Mode)", styles['Title']))
    elements.append(Spacer(1, 12))
    
    # FA Marks Report Table
    elements.append(Paragraph("FA Marks Report", styles['Heading2']))
    data = [report_df.columns.to_list()] + report_df.values.tolist()
    t = Table(data)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(t)
    elements.append(Spacer(1, 24))
    
    # Attendance Range Report Table
    elements.append(Paragraph("Attendance Range Report", styles['Heading2']))
    data1 = [report1_df.columns.to_list()] + report1_df.values.tolist()
    t1 = Table(data1)
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    elements.append(t1)
    
    doc.build(elements)
    return buffer.getvalue()

def create_docx(sheet_name, mode, report_df, report1_df):
    doc = Document()
    doc.add_heading(f'Student FA Marks Report - {sheet_name}', 0)
    doc.add_paragraph(f'Mode: {mode}')
    
    doc.add_heading('FA Marks Report', level=1)
    table = doc.add_table(rows=1, cols=len(report_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(report_df.columns):
        hdr_cells[i].text = str(col)
    for index, row in report_df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
            
    doc.add_page_break()
    doc.add_heading('Attendance Range Report', level=1)
    table1 = doc.add_table(rows=1, cols=len(report1_df.columns))
    table1.style = 'Table Grid'
    hdr_cells1 = table1.rows[0].cells
    for i, col in enumerate(report1_df.columns):
        hdr_cells1[i].text = str(col)
    for index, row in report1_df.iterrows():
        row_cells1 = table1.add_row().cells
        for i, val in enumerate(row):
            row_cells1[i].text = str(val)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def process_sheet(sheet_name, df, mode, original_filename):
    # Strip extension from original filename
    base_name = original_filename.rsplit('.', 1)[0] if '.' in original_filename else original_filename
    st.write(f"### Processing Sheet: {sheet_name} (from {original_filename})")

    # Required base columns
    base_columns = ["S.NO.", "REGD.", "CGPA", "Attendance"]
    missing_base = [c for c in base_columns if c not in df.columns]
    if missing_base:
        st.error(f"The sheet '{sheet_name}' must contain the required base columns: {', '.join(missing_base)}")
        return

    # Define test column sets
    test_1_4 = ["Test 1", "Test 2", "Test 3", "Test 4"]
    t_1_4 = ["T1", "T2", "T3", "T4"]
    
    test_columns = []
    
    if mode == 'Standard':
        has_test_1_4 = all(c in df.columns for c in test_1_4)
        has_t_1_4 = all(c in df.columns for c in t_1_4)

        if has_test_1_4 and has_t_1_4:
            st.error(f"❌ **Conflict in Standard Mode:** Sheet '{sheet_name}' contains both `{', '.join(test_1_4)}` AND `{', '.join(t_1_4)}`. Please use **Extended Mode** to process all 8 columns, or remove one set for Standard Mode.")
            return
        elif has_test_1_4:
            test_columns = test_1_4
        elif has_t_1_4:
            test_columns = t_1_4
        else:
            st.error(f"❌ **Standard Mode Error:** Sheet '{sheet_name}' must have either {', '.join(test_1_4)} OR {', '.join(t_1_4)}.")
            return
    else: # Extended Mode
        # In Extended Mode, strictly require all 8 columns
        required_extended = test_1_4 + t_1_4
        missing_extended = [c for c in required_extended if c not in df.columns]
        
        if missing_extended:
            st.error(f"❌ **Extended Mode Error:** Sheet '{sheet_name}' is missing columns: {', '.join(missing_extended)}. Only files with all 8 test columns are processed in this mode.")
            return
        
        test_columns = required_extended

    st.success(f"✅ Processing columns: {', '.join(test_columns)}")

    # Drop duplicates based on student registration number
    df = df.drop_duplicates(subset="REGD.")

    # Convert attendance to percentage if needed
    df["Attendance"] = df["Attendance"].apply(lambda x: float(str(x).strip('%')) if isinstance(x, str) and x.endswith('%') else x)

    # Add a category column
    df['Attendance_Percents'] = df["Attendance"].apply(lambda x: x * 100)
    df["Group"] = df["Attendance"].apply(categorize_attendance)

    # Convert test marks to integers, handling missing values
    for col in test_columns:
        df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0).astype(int)
        df[col] = df[col].replace(-1, 0).apply(lambda x: int(x))

    # Calculate total marks
    if mode == 'Standard':
        df["Total Marks"] = df[test_columns].sum(axis=1) * (3/4)
    else: # Extended Mode
        test_1_4 = ["Test 1", "Test 2", "Test 3", "Test 4"]
        t_1_4 = ["T1", "T2", "T3", "T4"]
        df["Total Marks"] = ((df[test_1_4].sum(axis=1) * (3/4)) + df[t_1_4].sum(axis=1)) / 2
    
    # Ensure marks are integers for categorization
    df["Total Marks"] = df["Total Marks"].astype(int)

    # Display the entire data with total marks
    st.write("### Student Data with Total Marks")
    st.dataframe(df)

    # Create a report structure
    attendance_ranges = ["≥75%", "≥65% - <75%", "≥50% - <65%", "<50%"]
    columns = ["<=10", ">10 & <=20", ">20 & <=30", ">30 & <=40", ">40 & <=50",">50 & <=60", "Total"]

    # Initialize report data with zeros
    report_data = {col: [0] * len(attendance_ranges) for col in columns}

    # Count students based on attendance and marks
    for index, row in df.iterrows():
        group = row["Group"]
        total_marks = row["Total Marks"]

        if total_marks <= 10:
            column = "<=10"
        elif 10 < total_marks <= 20:
            column = ">10 & <=20"
        elif 20 < total_marks <= 30:
            column = ">20 & <=30"
        elif 30 < total_marks <= 40:
            column = ">30 & <=40"
        elif 40 < total_marks <= 50:
            column = ">40 & <=50"
        elif 50 < total_marks <= 60:
            column = ">50 & <=60"
        else:
            continue

        report_data[column][attendance_ranges.index(group)] += 1
        report_data["Total"][attendance_ranges.index(group)] += 1

    # Calculate the total row
    report_df = pd.DataFrame(report_data, index=attendance_ranges)
    report_df.loc['Total'] = [sum(report_df[col]) for col in columns]

    report_df.reset_index(inplace=True)
    report_df.rename(columns={'index': 'Attendance'}, inplace=True)

    # Display the FA Marks Report
    st.write("### FA Marks Report")
    st.table(report_df)

    # Add attendance category
    df["Attendance_Category"] = df["Attendance_Percents"].apply(categorize_attendance_per)

    # Count students in each category
    attendance_counts = df["Attendance_Category"].value_counts().reindex([
        "0–49%", "50–64%", "65–74%", "75% and above"
    ], fill_value=0)

    # Display the report as a table
    st.write("### Attendance Range Report")
    report1_df = pd.DataFrame({
        "Attendance Range": attendance_counts.index,
        "No. of Students": attendance_counts.values
    })
    st.table(report1_df)

    # Export Options
    st.write("### 📤 Export Report")
    export_col1, export_col2 = st.columns([1, 2])
    
    with export_col1:
        export_format = st.selectbox("Select Format", ["Excel", "PDF", "Word"], key=f"fmt_{sheet_name}")
    
    with export_col2:
        if export_format == "Excel":
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                report_df.to_excel(writer, index=False, sheet_name='FA Marks Report')
                report1_df.to_excel(writer, index=False, sheet_name='Attendance Summary')
            data = output.getvalue()
            file_ext = "xlsx"
            mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        
        elif export_format == "PDF":
            data = create_pdf(sheet_name, st.session_state.mode, report_df, report1_df)
            file_ext = "pdf"
            mime = "application/pdf"
            
        else: # Word
            data = create_docx(sheet_name, st.session_state.mode, report_df, report1_df)
            file_ext = "docx"
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        st.download_button(
            label=f"Download {sheet_name} as {export_format}",
            data=data,
            file_name=f"{base_name}_{sheet_name}_{st.session_state.mode}.{file_ext}",
            mime=mime,
            key=f"dl_{sheet_name}_{export_format}",
            use_container_width=True
        )

def main():
    st.set_page_config(page_title="Student FA Marks Report", layout="wide")
    st.title("📊 Student FA Marks Report")

    # Mode Selection Buttons at the top
    if 'mode' not in st.session_state:
        st.session_state.mode = 'Standard'

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Standard Mode (Test 1, 2, 3, 4)", use_container_width=True, type="primary" if st.session_state.mode == 'Standard' else "secondary"):
            st.session_state.mode = 'Standard'
            st.rerun()
    with col2:
        if st.button("Extended Mode (+ T1, T2, T3, T4)", use_container_width=True, type="primary" if st.session_state.mode == 'Extended' else "secondary"):
            st.session_state.mode = 'Extended'
            st.rerun()

    # Define test columns based on mode
    if st.session_state.mode == 'Standard':
        test_cols = ["Test 1", "Test 2", "Test 3", "Test 4"]
        st.info("Current Mode: **Standard** (Processing Test 1 to Test 4)")
    else:
        test_cols = ["Test 1", "Test 2", "Test 3", "Test 4", "T1", "T2", "T3", "T4"]
        st.info("Current Mode: **Extended** (Processing Test 1-4 AND T1-4)")

    # Add instructions for the required Excel format
    st.markdown(
        f"""
        **📂 Excel File Requirements for {st.session_state.mode} Mode:**  
        The uploaded file must contain the following columns:  
        - `S.NO.`, `REGD.`, `CGPA`, `Attendance`
        - {", ".join([f"`{c}`" for c in test_cols])}
        - ✅ **Supports multiple sheets**: The application will process and display reports for each sheet.
        """
    )

    # Use a dynamic key for file uploader to clear files when mode changes
    uploader_key = f"file_uploader_{st.session_state.mode}"
    uploaded_file = st.file_uploader("Upload the Excel file", type=["xlsx"], key=uploader_key)

    if uploaded_file is not None:
        xls = pd.read_excel(uploaded_file, sheet_name=None)
        for sheet_name, df in xls.items():
            process_sheet(sheet_name, df, st.session_state.mode, uploaded_file.name)

if __name__ == "__main__":
    main()
